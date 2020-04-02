import urllib.request
from docx import Document
from docx.shared import Inches

document = Document()

#TIMETABLE BEGIN

timetable = [
    {
        'Monday': 'Rom', 
        'Tuesday': 'Bio',
        'Wednesday': 'Rom',             # FIRST PERIOD
        'Thursday': 'Log',
        'Friday': 'Ch',
    },
    {
        'Monday': 'Mat', 
        'Tuesday': 'Fr',
        'Wednesday': 'Mat',             # SECOND PERIOD
        'Thursday': 'Inf',
        'Friday': 'Inf',
    },
    {
        'Monday': 'Geo', 
        'Tuesday': 'Ist',
        'Wednesday': 'Mat',             # THIRD PERIDO
        'Thursday': 'Bio',
        'Friday': 'ITC',
    },
    {
        'Monday': 'Rom', 
        'Tuesday': 'Mat',
        'Wednesday': 'Eng',             # FOURTH PERIOD
        'Thursday': 'Mat',
        'Friday': 'Arts' ,
    },
    {
        'Monday': 'Eng', 
        'Tuesday': 'Fiz',
        'Wednesday': 'Rom',             # FIFTH PERIOD
        'Thursday': 'Ch',
        'Friday': 'Fiz',
    },
    {
        'Monday': 'Drg', 
        'Tuesday': 'Free',
        'Wednesday': 'Fiz',             # SIXTH PERIOD
        'Thursday': 'Fr',
        'Friday': 'Ed-Fiz',
    },
    {
        'Monday': 'Free', 
        'Tuesday': 'Free',
        'Wednesday': 'Free',             # SEVENTH PERIOD
        'Thursday': 'Rel',
        'Friday': 'Free',
    }
]

#TIMETABLE END

document.add_heading('Colegiul National Costache Negruzzi', 0) 
document.add_heading('Iasi, Romania', level=1) 
document.add_heading('Proiect realizat de elevii clasei 9A', level=2)
document.add_heading('In coordonare cu Ms. Grumeza Diana', level=3)
#HEADINGS
urllib.request.urlretrieve("https://i.ytimg.com/vi/ikEYqV0y-SE/maxresdefault.jpg", "maxresdefault.jpg") 
document.add_picture('maxresdefault.jpg', width=Inches(5)) 
#PICTURE
document.add_paragraph('our best teachers').bold = True 
document.add_paragraph('grum', style='ListBullet')
document.add_paragraph('foca', style='ListBullet')
document.add_paragraph('soro', style='ListBullet')
document.add_paragraph('our best results').bold = True 
document.add_paragraph('geo', style='ListNumber')
document.add_paragraph('info', style='ListNumber')
document.add_paragraph('rel', style='ListNumber')
#OL/UL

document.add_page_break()

#TIMETABLE PRINT BEGIN
document.add_heading('9A\'s timetable', level=2)
table = document.add_table(rows=1, cols=5)
cells = table.rows[0].cells
cells[0].text = 'Mon'
cells[1].text = 'Tue'
cells[2].text = 'Wed'
cells[3].text = 'Thu'
cells[4].text = 'Fri'

for item in timetable:
    row_cells = table.add_row().cells
    row_cells[0].text = item['Monday']
    row_cells[1].text = item['Tuesday']
    row_cells[2].text = item['Wednesday']
    row_cells[3].text = item['Thursday']
    row_cells[4].text = item['Friday']

#TIMETABLE PRINT END

document.add_page_break()

document.save('test.docx')




